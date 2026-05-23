"""사용자 업로드 문서 → 마크다운 추출.

지원 포맷:
- DOCX: python-docx — 단락·헤딩·표·리스트 보존
- HWPX: python-hwpx — 한컴 신규 포맷 (XML 기반)
- HWP: 레거시 바이너리. python-hwp/olefile 폴백 — 텍스트만 추출 가능
- PDF: pypdf — 텍스트 기반 PDF만. 스캔 PDF는 OCR 필요(미지원)

설계 원칙:
- IR을 거치지 않고 곧장 마크다운으로 — 사용자가 에디터에서 다시 편집 가능
- 헤딩·표·글머리는 마크다운 구조로 변환 (변환 파이프라인이 다시 IR로 파싱)
- 실패는 명확한 사용자 메시지 반환
"""
from __future__ import annotations

import re
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from app.core.logging import get_logger

log = get_logger(__name__)


@dataclass
class ExtractResult:
    markdown: str
    title: str = ""
    warnings: list[str] = None  # type: ignore[assignment]
    format: str = ""

    def __post_init__(self) -> None:
        if self.warnings is None:
            self.warnings = []


class UnsupportedFormatError(ValueError):
    pass


# ─────────────────────────────────────────────────────────────────────────────
# DOCX
# ─────────────────────────────────────────────────────────────────────────────

def _docx_to_markdown(content: bytes) -> ExtractResult:
    from docx import Document
    from docx.table import Table as DocxTable

    doc = Document(BytesIO(content))
    lines: list[str] = []
    title = ""
    warnings: list[str] = []

    def _para_to_md(para) -> str:
        # 스타일 기반 헤딩 감지
        style_name = (para.style.name or "").lower() if para.style else ""
        text = para.text.rstrip()
        if not text:
            return ""
        if style_name.startswith("heading 1") or style_name == "title":
            return f"# {text}"
        if style_name.startswith("heading 2"):
            return f"## {text}"
        if style_name.startswith("heading 3"):
            return f"### {text}"
        if style_name.startswith("heading 4"):
            return f"#### {text}"
        if style_name.startswith("heading"):
            return f"##### {text}"
        # 한국 공문 글머리 자동 감지 (□ ○ ― ※)
        m = re.match(r"^([□○―※])\s*(.+)", text)
        if m:
            bullet, body = m.groups()
            depth = {"□": 0, "○": 1, "―": 2, "※": 3}.get(bullet, 0)
            return "  " * depth + f"- {body}"
        # 리스트 스타일
        if style_name.startswith("list bullet") or style_name.startswith("list paragraph"):
            return f"- {text}"
        return text

    def _table_to_md(table: DocxTable) -> list[str]:
        rows = []
        for row in table.rows:
            cells = [c.text.replace("\n", " ").replace("|", "\\|").strip() for c in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        if not rows:
            return []
        cols = len(table.columns)
        sep = "| " + " | ".join(["---"] * cols) + " |"
        rows.insert(1, sep)
        return rows + [""]

    # body 순회 — 단락과 표 순서 유지
    # python-docx의 element 자식을 순회해야 순서 보존됨
    from docx.document import Document as DocClass
    if isinstance(doc, DocClass):
        body = doc.element.body
        for child in body.iterchildren():
            tag = child.tag.split("}")[-1]  # remove namespace
            if tag == "p":
                # paragraph element
                for para in doc.paragraphs:
                    if para._element is child:
                        md = _para_to_md(para)
                        if md:
                            # 첫 번째 H1을 title로
                            if not title and md.startswith("# "):
                                title = md[2:].strip()
                            lines.append(md)
                        break
            elif tag == "tbl":
                for tbl in doc.tables:
                    if tbl._element is child:
                        lines.extend(_table_to_md(tbl))
                        break

    # fallback: 위 순회가 실패해도 단락만이라도
    if not lines:
        for para in doc.paragraphs:
            md = _para_to_md(para)
            if md:
                if not title and md.startswith("# "):
                    title = md[2:].strip()
                lines.append(md)
        for tbl in doc.tables:
            lines.extend(_table_to_md(tbl))

    md_text = "\n\n".join(lines)
    if not title and lines:
        title = "업로드된 DOCX 문서"
    return ExtractResult(markdown=md_text, title=title, warnings=warnings, format="docx")


# ─────────────────────────────────────────────────────────────────────────────
# HWPX (신규 한컴 포맷)
# ─────────────────────────────────────────────────────────────────────────────

def _hwpx_to_markdown(content: bytes) -> ExtractResult:
    import tempfile

    import hwpx  # type: ignore[import-not-found]

    warnings: list[str] = []
    title = ""

    # python-hwpx는 파일 경로를 요구. 임시 파일로 우회.
    with tempfile.NamedTemporaryFile(suffix=".hwpx", delete=False) as tf:
        tf.write(content)
        tmp_path = Path(tf.name)
    try:
        doc = hwpx.HwpxDocument.open(str(tmp_path))

        # TextExtractor가 있으면 활용
        try:
            extractor = hwpx.TextExtractor(doc)  # type: ignore[attr-defined]
            text = extractor.extract()
        except Exception:
            # 폴백: paragraphs 순회
            text_parts: list[str] = []
            for para in getattr(doc, "paragraphs", []):
                try:
                    if hasattr(para, "text"):
                        text_parts.append(str(para.text))
                    elif hasattr(para, "get_text"):
                        text_parts.append(str(para.get_text()))
                except Exception:
                    continue
            text = "\n\n".join(text_parts)

        if not text.strip():
            warnings.append("본문 텍스트를 추출하지 못했습니다.")

        # 줄 단위로 한국 공문 글머리 자동 감지 + 마크다운화
        lines = []
        for line in text.splitlines():
            stripped = line.strip()
            if not stripped:
                lines.append("")
                continue
            # 첫 의미있는 줄을 제목으로
            if not title:
                title = stripped[:100]
            m = re.match(r"^([□○―※])\s*(.+)", stripped)
            if m:
                bullet, body = m.groups()
                depth = {"□": 0, "○": 1, "―": 2, "※": 3}.get(bullet, 0)
                lines.append("  " * depth + f"- {body}")
            else:
                lines.append(stripped)

        md = "\n".join(lines)
        return ExtractResult(markdown=md, title=title, warnings=warnings, format="hwpx")
    finally:
        tmp_path.unlink(missing_ok=True)


# ─────────────────────────────────────────────────────────────────────────────
# HWP (레거시 바이너리)
# ─────────────────────────────────────────────────────────────────────────────

def _hwp_to_markdown(content: bytes) -> ExtractResult:
    """HWP 5.x 레거시 (.hwp) 파싱.

    우선순위:
    1. pyhwp의 hwp5txt CLI — 본문 텍스트 완전 추출
    2. olefile PrvText 스트림 — 미리보기 텍스트만 (폴백)

    HWP 5.0 스펙 호환성 노트:
    - 한컴이 공개한 HWP 5.0 스펙 문서에는 실제 바이너리와의 불일치 27건이 있음
      (BorderFill 직렬화 순서, GShapeObject 4바이트 오프셋, ParagraphShape 미사용 필드,
       PrvImage PNG 포맷 등). rhwp 프로젝트의 hwp_spec_errata 공개 자료 참고.
    - 본 DocuAX는 pyhwp(올레 기반 + 자체 errata 처리)에 의존.
    - 한국 저작권법 §101조의4 / 미국 DMCA §1201(f) / EU 소프트웨어 지침 §6 모두
      상호운용성을 위한 역공학을 허용. 본 추출 기능은 사용자가 본인 소유 파일을
      DocuAX에서 활용하기 위한 합법적 호환성 처리임.
    - "한컴", "HWP", "HWPX"는 (주)한글과컴퓨터의 등록 상표. DocuAX는 한컴과
      제휴·후원·승인 관계 없는 독립 프로젝트.
    """
    warnings: list[str] = []
    extracted = ""
    title = "업로드된 HWP 문서"

    # 1) pyhwp CLI 경로 (있으면 본문 추출)
    cli_text = _try_pyhwp_cli(content)
    if cli_text:
        extracted = cli_text
        # 첫 의미있는 줄을 제목으로
        for line in extracted.splitlines():
            s = line.strip()
            if s:
                title = s[:100]
                break
        # 한국 공문 글머리 자동 마크다운화
        lines = []
        for raw in extracted.splitlines():
            s = raw.strip()
            if not s:
                lines.append("")
                continue
            m = re.match(r"^([□○―※])\s*(.+)", s)
            if m:
                bullet, body = m.groups()
                depth = {"□": 0, "○": 1, "―": 2, "※": 3}.get(bullet, 0)
                lines.append("  " * depth + f"- {body}")
            else:
                lines.append(s)
        extracted = "\n".join(lines)
        return ExtractResult(markdown=extracted, title=title, warnings=warnings, format="hwp")

    # 2) olefile PrvText 폴백
    try:
        import olefile
    except ImportError:
        return ExtractResult(
            markdown="",
            title=title,
            warnings=[
                "HWP 파싱 라이브러리가 사용 불가합니다 (pyhwp/olefile 미설치).",
                "권장: 한컴 한글에서 다른 이름으로 저장 → HWPX 포맷으로 변환 후 재업로드.",
            ],
            format="hwp",
        )

    try:
        ole = olefile.OleFileIO(BytesIO(content))
        if ole.exists("PrvText"):
            with ole.openstream("PrvText") as f:
                raw = f.read()
            try:
                extracted = raw.decode("utf-16-le", errors="ignore").strip("\x00")
            except UnicodeDecodeError:
                extracted = raw.decode("utf-8", errors="ignore")
            warnings.append(
                "pyhwp CLI 추출이 실패해 미리보기 텍스트만 추출되었습니다. 전체 본문이 필요하면 HWPX로 변환 후 재업로드하세요."
            )
        else:
            warnings.append("HWP 파일에서 텍스트를 추출할 수 없습니다. HWPX로 변환 후 재업로드해주세요.")
        ole.close()
    except Exception as e:  # noqa: BLE001
        log.warning("HWP 추출 실패", error=str(e))
        warnings.append(f"HWP 파싱 실패: {e}. HWPX로 변환 후 재업로드해주세요.")

    return ExtractResult(markdown=extracted, title=title, warnings=warnings, format="hwp")


def _try_pyhwp_cli(content: bytes) -> str:
    """pyhwp의 hwp5txt CLI로 본문 텍스트 추출. 미설치/실패 시 빈 문자열."""
    import shutil
    import subprocess
    import tempfile

    cli = shutil.which("hwp5txt")
    if not cli:
        return ""

    with tempfile.NamedTemporaryFile(suffix=".hwp", delete=False) as tf:
        tf.write(content)
        tmp_path = tf.name
    try:
        result = subprocess.run(
            [cli, tmp_path],
            check=False, capture_output=True, timeout=60,
        )
        if result.returncode != 0:
            log.warning("hwp5txt 실패", stderr=result.stderr[:200].decode(errors="replace"))
            return ""
        # hwp5txt는 UTF-8 stdout
        try:
            return result.stdout.decode("utf-8", errors="replace").strip()
        except Exception:  # noqa: BLE001
            return ""
    except (subprocess.TimeoutExpired, OSError) as e:
        log.warning("hwp5txt 호출 실패", error=str(e))
        return ""
    finally:
        from pathlib import Path as _P
        _P(tmp_path).unlink(missing_ok=True)


# ─────────────────────────────────────────────────────────────────────────────
# PDF
# ─────────────────────────────────────────────────────────────────────────────

def _ocr_pdf_fallback(content: bytes) -> tuple[str, list[str]]:
    """OCR provider가 활성화돼 있으면 PDF 전체 OCR. 비활성/실패 시 빈 텍스트."""
    try:
        from app.providers.ocr import get_ocr_provider
        provider = get_ocr_provider()
    except Exception as e:  # noqa: BLE001
        return "", [f"OCR provider 초기화 실패: {e}"]

    if provider.name == "none":
        return "", ["스캔 PDF로 보입니다. 설정에서 OCR provider(Tesseract/CLOVA)를 활성화하세요."]

    try:
        import asyncio
        result = asyncio.run(provider.ocr_pdf(content))
    except Exception as e:  # noqa: BLE001
        return "", [f"OCR 실패 ({provider.name}): {e}"]

    if not result.text.strip():
        return "", [f"OCR이 텍스트를 추출하지 못했습니다 ({provider.name})"]

    warns = [
        f"OCR({result.provider})로 추출됨 · {result.total_chars}자 · 페이지 {len(result.pages)}"
    ]
    for p in result.pages:
        warns.extend(p.warnings)
    return result.text, warns


def _pdf_to_markdown(content: bytes, *, allow_ocr: bool = True) -> ExtractResult:
    from pypdf import PdfReader

    warnings: list[str] = []
    reader = PdfReader(BytesIO(content))
    title = (reader.metadata.title if reader.metadata else None) or "업로드된 PDF 문서"

    pages_text: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
        except Exception as e:  # noqa: BLE001
            warnings.append(f"{i + 1}쪽 추출 실패: {e}")
            continue
        if text.strip():
            pages_text.append(text)

    # 텍스트 추출 0건 → OCR 폴백 시도
    if not pages_text and allow_ocr:
        ocr_text, ocr_warnings = _ocr_pdf_fallback(content)
        warnings.extend(ocr_warnings)
        if ocr_text:
            pages_text = [ocr_text]

    if not pages_text:
        warnings.append(
            "텍스트를 추출할 수 없습니다. 스캔 PDF인 경우 OCR provider(Tesseract/CLOVA) 설정 후 재시도하세요."
        )
        return ExtractResult(markdown="", title=title, warnings=warnings, format="pdf")

    # 페이지 단위로 합치고, 짧은 줄이 연속이면 합치는 휴리스틱
    lines: list[str] = []
    for page in pages_text:
        for raw in page.splitlines():
            stripped = raw.strip()
            if not stripped:
                lines.append("")
                continue
            # 한국 공문 글머리 자동 감지
            m = re.match(r"^([□○―※])\s*(.+)", stripped)
            if m:
                bullet, body = m.groups()
                depth = {"□": 0, "○": 1, "―": 2, "※": 3}.get(bullet, 0)
                lines.append("  " * depth + f"- {body}")
            else:
                lines.append(stripped)
        lines.append("")  # 페이지 구분

    md = "\n".join(lines)
    return ExtractResult(markdown=md, title=title, warnings=warnings, format="pdf")


# ─────────────────────────────────────────────────────────────────────────────
# HTML
# ─────────────────────────────────────────────────────────────────────────────

def _html_to_markdown(content: bytes) -> ExtractResult:
    """HTML → 마크다운. 강의 자료·웹 페이지·노션 export 등 활용.

    - <h1>~<h6> 헤딩 보존
    - <table>·<ul>·<ol>·<blockquote>·<code> 보존
    - <style>·<script> 제거 (디자인 CSS는 변환 후 양식 매핑에서 재구성)
    - 한국 공문 글머리(□ ○ ― ※)는 자동 인식
    """
    try:
        from bs4 import BeautifulSoup
        from markdownify import markdownify
    except ImportError as e:
        return ExtractResult(
            markdown="",
            title="HTML 문서",
            warnings=[f"HTML 라이브러리 미설치: {e}. pip install beautifulsoup4 markdownify"],
            format="html",
        )

    warnings: list[str] = []
    # 인코딩 자동 감지 (UTF-8 기본, 한글 사이트 EUC-KR도 시도)
    try:
        text = content.decode("utf-8")
    except UnicodeDecodeError:
        try:
            text = content.decode("euc-kr")
            warnings.append("EUC-KR 인코딩으로 디코딩됨")
        except UnicodeDecodeError:
            text = content.decode("utf-8", errors="replace")
            warnings.append("일부 문자 손실")

    soup = BeautifulSoup(text, "html.parser")
    # title
    title = ""
    if soup.title and soup.title.string:
        title = soup.title.string.strip()
    elif soup.find("h1"):
        title = soup.find("h1").get_text(strip=True)
    else:
        title = "업로드된 HTML 문서"

    # 노이즈 제거 — <style>, <script>, <meta>, <link>, <noscript>, <iframe>
    for tag in soup(["style", "script", "meta", "link", "noscript", "iframe", "svg"]):
        tag.decompose()

    # body가 있으면 본문만, 없으면 전체
    target = soup.body if soup.body else soup

    # markdownify로 변환
    md_text = markdownify(
        str(target),
        heading_style="ATX",  # # ## 형식
        bullets="-",
        strong_em_symbol="*",
        strip=["a"],  # 링크는 텍스트만
    )

    # 빈 줄 정리 — 연속 빈 줄 3개 이상은 2개로
    lines = md_text.splitlines()
    cleaned: list[str] = []
    empty_count = 0
    for line in lines:
        if not line.strip():
            empty_count += 1
            if empty_count <= 2:
                cleaned.append("")
        else:
            cleaned.append(line.rstrip())
            empty_count = 0
    md_text = "\n".join(cleaned).strip()

    return ExtractResult(markdown=md_text, title=title, warnings=warnings, format="html")


# ─────────────────────────────────────────────────────────────────────────────
# 진입점
# ─────────────────────────────────────────────────────────────────────────────

def extract_document(filename: str, content: bytes) -> ExtractResult:
    """파일명 확장자로 포맷 판별 후 마크다운 추출."""
    suffix = Path(filename).suffix.lower()
    try:
        if suffix == ".docx":
            return _docx_to_markdown(content)
        if suffix == ".hwpx":
            return _hwpx_to_markdown(content)
        if suffix == ".hwp":
            return _hwp_to_markdown(content)
        if suffix in (".html", ".htm"):
            return _html_to_markdown(content)
        if suffix == ".pdf":
            return _pdf_to_markdown(content)
        if suffix == ".md":
            # 이미 마크다운이면 그대로 통과 (LLM 결과 저장본 등)
            text = content.decode("utf-8", errors="replace")
            first_h1 = ""
            for line in text.splitlines():
                if line.startswith("# "):
                    first_h1 = line[2:].strip()
                    break
            return ExtractResult(markdown=text, title=first_h1 or filename, warnings=[], format="md")
    except UnsupportedFormatError:
        raise
    except Exception as e:  # noqa: BLE001
        log.exception("문서 추출 실패", filename=filename)
        return ExtractResult(
            markdown="",
            title=filename,
            warnings=[f"추출 실패: {e}"],
            format=suffix.lstrip("."),
        )

    raise UnsupportedFormatError(
        f"지원하지 않는 포맷: {suffix}. 지원: .docx, .hwpx, .hwp, .pdf, .html, .htm, .md"
    )
