"""DOCX 렌더러 — python-docx.

DocumentIR.blocks를 순회하면서 docx 문서로 직렬화.
한국 공문 4단계 글머리(□ ○ ― ※)와 표 서식을 정확히 표현.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

from app.core.logging import get_logger
from app.pipeline.ir import Block, BlockType, CoverData, DocumentIR, InlineRun, TableCell
from app.renderers.base import Renderer
from app.services.organization_profile import OrganizationProfile, profile_for_document
from app.services.visuals import (
    materialize_image,
    render_chart_to_png,
    render_diagram_to_png,
    render_equation_to_png,
    to_cm,
)
from app.services.visuals.cache import resolve_image_to_path

log = get_logger(__name__)


def _hex_to_rgb(c: str | None) -> RGBColor | None:
    if not c:
        return None
    s = c.lstrip("#")
    if len(s) != 6:
        return None
    try:
        return RGBColor(int(s[0:2], 16), int(s[2:4], 16), int(s[4:6], 16))
    except ValueError:
        return None


def _apply_run(docx_run, run: InlineRun) -> None:
    docx_run.text = run.text
    docx_run.bold = run.bold
    docx_run.italic = run.italic
    docx_run.underline = run.underline
    if run.strikethrough:
        docx_run.font.strike = True
    if run.superscript:
        docx_run.font.superscript = True
    if run.subscript:
        docx_run.font.subscript = True
    if run.font_family:
        docx_run.font.name = run.font_family
        # 동아시아 폰트도 동일하게 지정
        rpr = docx_run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            from docx.oxml import OxmlElement
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), run.font_family)
    if run.font_size:
        docx_run.font.size = Pt(run.font_size)
    if run.color:
        rgb = _hex_to_rgb(run.color)
        if rgb:
            docx_run.font.color.rgb = rgb


def _shade_cell(cell, color: str | None) -> None:
    if not color:
        return
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color.lstrip("#"))
    tc_pr.append(shd)


def _set_cell_borders(cell, style: str) -> None:
    """단일 셀 테두리 — none이면 제거."""
    if style != "none":
        return
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def _write_block(doc, blk: Block, profile: OrganizationProfile | None = None) -> None:
    if blk.type == BlockType.HEADING and blk.runs:
        level = max(0, min(9, blk.heading_level))
        p = doc.add_heading("", level=level)
        # 조직 프로파일 헤딩 색·서체 우선 적용
        heading_color = None
        heading_font = None
        heading_size = None
        if profile:
            heading_color = _hex_to_rgb(
                profile.brand_color_hex if level <= 2 else profile.accent_color_hex
            )
            heading_font = profile.font_korean_heading or profile.font_korean
            heading_size = {1: profile.h1_font_size_pt, 2: profile.h2_font_size_pt, 3: profile.h3_font_size_pt}.get(level)
        for r in blk.runs:
            run = p.add_run()
            _apply_run(run, r)
            if heading_color and not r.color:
                run.font.color.rgb = heading_color
            if heading_font and not r.font_family:
                run.font.name = heading_font
                rpr = run._element.get_or_add_rPr()
                rfonts = rpr.find(qn("w:rFonts"))
                if rfonts is None:
                    from docx.oxml import OxmlElement
                    rfonts = OxmlElement("w:rFonts")
                    rpr.append(rfonts)
                rfonts.set(qn("w:eastAsia"), heading_font)
            if heading_size and not r.font_size:
                run.font.size = Pt(heading_size)
        return

    if blk.type == BlockType.LIST_ITEM and blk.list_item:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5 * blk.list_item.depth)
        marker = (
            f"{blk.list_item.index}{blk.list_item.order_format.replace('1.', '.')} "
            if blk.list_item.ordered
            else f"{blk.list_item.bullet_marker} "
        )
        m_run = p.add_run(marker)
        m_run.bold = True
        for r in blk.list_item.runs:
            _apply_run(p.add_run(), r)
        return

    if blk.type == BlockType.QUOTE:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        for r in blk.runs:
            _apply_run(p.add_run(), r).italic = True  # type: ignore[union-attr]
        return

    if blk.type == BlockType.CODE:
        p = doc.add_paragraph()
        for r in blk.runs:
            docx_run = p.add_run(r.text)
            docx_run.font.name = "JetBrains Mono"
            docx_run.font.size = Pt(10)
        return

    if blk.type == BlockType.TABLE and blk.table:
        rows = blk.table.rows
        if not rows:
            return
        col_count = max(len(r) for r in rows)
        tbl = doc.add_table(rows=len(rows), cols=col_count)
        if blk.table.border_style != "none":
            tbl.style = "Table Grid"
        # 조직 프로파일: 헤더 행에 브랜드 배경 (없으면 기본 #F2F2F2)
        if profile and blk.table.header_row and rows:
            from docx.oxml import OxmlElement
            for ci in range(min(col_count, len(rows[0]))):
                try:
                    cell = tbl.rows[0].cells[ci]
                    tc_pr = cell._tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:fill"), profile.table_header_bg_hex.lstrip("#"))
                    tc_pr.append(shd)
                except Exception:
                    continue
        # 페이지 정렬
        align_map = {"left": WD_TABLE_ALIGNMENT.LEFT, "center": WD_TABLE_ALIGNMENT.CENTER, "right": WD_TABLE_ALIGNMENT.RIGHT}
        tbl.alignment = align_map.get(blk.table.align, WD_TABLE_ALIGNMENT.CENTER)

        for ri, row in enumerate(rows):
            for ci, cell_data in enumerate(row):
                if ci >= col_count:
                    break
                if cell_data.colspan == 0 or cell_data.rowspan == 0:
                    continue  # 병합된 영역 — 헤드 셀이 처리
                cell = tbl.rows[ri].cells[ci]
                # 셀 내용
                cell.text = ""  # 기본 단락 비움
                p = cell.paragraphs[0]
                align_p = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER, "right": WD_ALIGN_PARAGRAPH.RIGHT}
                p.alignment = align_p.get(cell_data.align, WD_ALIGN_PARAGRAPH.LEFT)
                for r in cell_data.runs:
                    _apply_run(p.add_run(), r)
                _shade_cell(cell, cell_data.background)
                _set_cell_borders(cell, cell_data.border)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                # 셀 병합
                if cell_data.rowspan > 1 or cell_data.colspan > 1:
                    end_r = min(len(rows) - 1, ri + cell_data.rowspan - 1)
                    end_c = min(col_count - 1, ci + cell_data.colspan - 1)
                    cell.merge(tbl.rows[end_r].cells[end_c])
        return

    if blk.type == BlockType.BOX:
        p = doc.add_paragraph()
        for r in blk.runs:
            _apply_run(p.add_run(), r)
        # 박스 — paragraph border (docx XML)
        from docx.oxml import OxmlElement
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        for edge in ("top", "left", "bottom", "right"):
            b = OxmlElement(f"w:{edge}")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), "8")
            b.set(qn("w:color"), "auto")
            pBdr.append(b)
        pPr.append(pBdr)
        return

    if blk.type == BlockType.THEMATIC_BREAK:
        doc.add_paragraph("─" * 50)
        return

    # ── 시각 요소 ──
    if blk.type == BlockType.IMAGE and blk.image:
        _write_image(doc, blk.image)
        return

    if blk.type == BlockType.DIAGRAM and blk.diagram:
        png = render_diagram_to_png(engine=blk.diagram.engine, source=blk.diagram.source)
        if png:
            cm = to_cm(blk.diagram.width)
            _embed_picture(doc, png, blk.diagram.caption, blk.diagram.align,
                           blk.diagram.width_pt, override_cm=cm)
        else:
            p = doc.add_paragraph()
            p.add_run(f"[다이어그램 렌더 실패: {blk.diagram.engine}]").italic = True
        return

    if blk.type == BlockType.CHART and blk.chart:
        png = render_chart_to_png(blk.chart.spec)
        if png:
            cm = to_cm(blk.chart.width)
            _embed_picture(doc, png, blk.chart.caption, blk.chart.align,
                           blk.chart.width_pt, override_cm=cm)
        else:
            p = doc.add_paragraph()
            p.add_run(f"[차트 렌더 실패]").italic = True
        return

    if blk.type == BlockType.EQUATION and blk.equation:
        png = render_equation_to_png(latex=blk.equation.latex, display=blk.equation.display)
        if png:
            cm = to_cm(blk.equation.width)
            _embed_picture(doc, png, "", blk.equation.align, None,
                           max_width_cm=10, override_cm=cm)
        else:
            p = doc.add_paragraph()
            p.add_run(blk.equation.latex).italic = True
        return

    # 기본: 단락
    p = doc.add_paragraph()
    for r in blk.runs:
        _apply_run(p.add_run(), r)


def _write_image(doc, img) -> None:
    """이미지 블록 — 파일을 materialize 한 뒤 임베드."""
    path = resolve_image_to_path(img)
    if not path:
        p = doc.add_paragraph()
        p.add_run(f"[이미지 로드 실패: {img.alt or img.src}]").italic = True
        return
    override_cm = to_cm(img.width)
    _embed_picture(doc, path, img.caption, img.align, img.width_pt, override_cm=override_cm)


def _embed_picture(
    doc,
    image_path: Path,
    caption: str = "",
    align: str = "center",
    width_pt: float | None = None,
    max_width_cm: float = 16.0,
    override_cm: float | None = None,
) -> None:
    """python-docx 로 그림 임베드 + 캡션 추가.

    크기 우선순위:
      1. override_cm — width 문자열 ("70%", "120mm" 등) 변환 결과
      2. width_pt    — 매크로/이전 값
      3. 페이지 폭의 80% 또는 max_width_cm
    """
    p = doc.add_paragraph()
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    p.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    try:
        if override_cm is not None and override_cm > 0:
            # 페이지 폭 16cm 한계 적용
            cm = min(override_cm, max_width_cm)
            p.add_run().add_picture(str(image_path), width=Cm(cm))
        elif width_pt:
            p.add_run().add_picture(str(image_path), width=Pt(width_pt))
        else:
            p.add_run().add_picture(str(image_path), width=Cm(min(max_width_cm, 15.0)))
    except Exception as e:  # noqa: BLE001
        log.warning("DOCX 이미지 임베드 실패", path=str(image_path), error=str(e))
        p.add_run(f"[이미지 로드 실패: {image_path.name}]").italic = True
        return

    if caption:
        cap = doc.add_paragraph()
        cap.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.CENTER)
        cap_run = cap.add_run(caption)
        cap_run.italic = True
        cap_run.font.size = Pt(9)
        cap_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def _write_cover_page(doc, cover: CoverData, profile: OrganizationProfile | None) -> None:
    """표지 페이지 렌더 — 단독 페이지."""
    brand_rgb = _hex_to_rgb(profile.brand_color_hex) if profile else _hex_to_rgb("#1E2761")

    # 상단 여백 — 빈 단락 6개로 시각 중앙 정렬에 가깝게
    for _ in range(4):
        doc.add_paragraph("")

    # 로고 (있으면)
    if cover.logo_path:
        try:
            logo_path = materialize_image(local_path=cover.logo_path, src=cover.logo_path)
            if logo_path:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(str(logo_path), width=Cm(4.0))
                doc.add_paragraph("")
        except Exception:
            pass

    # 기관명
    if cover.organization:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(cover.organization)
        r.font.size = Pt(14)
        r.bold = True
        if brand_rgb:
            r.font.color.rgb = brand_rgb

    if cover.department:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(cover.department)
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # 분류 (대외비 등)
    if cover.classification:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"[ {cover.classification} ]")
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        r.bold = True

    # 큰 빈 줄
    for _ in range(3):
        doc.add_paragraph("")

    # 제목 (큰 글씨)
    if cover.title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(cover.title)
        r.font.size = Pt(28)
        r.bold = True
        if brand_rgb:
            r.font.color.rgb = brand_rgb

    if cover.subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(cover.subtitle)
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # 큰 빈 줄
    for _ in range(8):
        doc.add_paragraph("")

    # 하단 작성자 · 날짜 · 문서번호
    info_lines: list[str] = []
    if cover.author:
        info_lines.append(f"작 성 자 :  {cover.author}")
    if cover.date:
        info_lines.append(f"작성일자 :  {cover.date}")
    if cover.document_number:
        info_lines.append(f"문서번호 :  {cover.document_number}")
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(11)

    # 페이지 분리
    page_break = doc.add_paragraph()
    page_break.add_run().add_break(WD_BREAK.PAGE)


def _apply_profile_setup(doc, profile: OrganizationProfile) -> None:
    """페이지 여백·기본 서체·헤더/푸터를 조직 프로파일대로 설정."""
    from docx.shared import Mm
    # 1) 페이지 여백
    try:
        for section in doc.sections:
            section.top_margin = Mm(profile.margin_top_mm)
            section.bottom_margin = Mm(profile.margin_bottom_mm)
            section.left_margin = Mm(profile.margin_left_mm)
            section.right_margin = Mm(profile.margin_right_mm)
    except Exception:
        pass

    # 2) 기본 본문 서체 — 'Normal' 스타일에 적용
    try:
        normal = doc.styles["Normal"]
        normal.font.name = profile.font_korean
        normal.font.size = Pt(profile.body_font_size_pt)
        # 동아시아 폰트
        from docx.oxml import OxmlElement
        rpr = normal.element.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            normal.element.insert(0, rpr)
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), profile.font_korean)
        rfonts.set(qn("w:ascii"), profile.font_korean)
        rfonts.set(qn("w:hAnsi"), profile.font_korean)
    except Exception:
        pass

    # 3) 헤더/푸터
    try:
        for section in doc.sections:
            if profile.header_text:
                section.header.paragraphs[0].text = profile.header_text
            if profile.footer_text:
                section.footer.paragraphs[0].text = profile.footer_text
    except Exception:
        pass


class DocxRenderer(Renderer):
    extension = ".docx"
    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def render(self, ir: DocumentIR, output_path: Path) -> Path:
        doc = Document()
        # 조직 프로파일 — convert 시 remember_for_document 로 미리 캐싱됨
        profile = profile_for_document(ir.document_id)
        if profile:
            _apply_profile_setup(doc, profile)

        # 표지(있으면) — 본문 위 별도 페이지
        if ir.cover:
            try:
                _write_cover_page(doc, ir.cover, profile)
            except Exception as e:  # noqa: BLE001
                log.warning("DOCX 표지 렌더 실패", error=str(e))

        # 제목 (표지에 이미 들어갔으면 중복 방지)
        if ir.title and (not ir.cover or not ir.cover.title):
            t = doc.add_heading(ir.title, level=0)
            t.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if profile:
                brand_rgb = _hex_to_rgb(profile.brand_color_hex)
                for run in t.runs:
                    if brand_rgb:
                        run.font.color.rgb = brand_rgb
                    if profile.font_korean_heading:
                        run.font.name = profile.font_korean_heading
        for blk in ir.blocks:
            try:
                _write_block(doc, blk, profile=profile)
            except Exception as e:  # noqa: BLE001
                # 한 블록 실패해도 전체 변환을 막지 않음
                log.warning("DOCX 블록 실패", block=blk.id, type=blk.type, error=str(e))
                continue
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return output_path
